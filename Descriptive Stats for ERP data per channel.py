"""
Descriptive Stats for ERP data per channel

1. Input:
   - The script reads data from Excel files stored in subdirectories corresponding to different experimental conditions (e.g., 'keypress', 'robot_sync', 'robot_async') and events (e.g., '64', '65').
   - Each Excel file should follow the naming convention `average_erp_data_<condition>_event_<event>.xlsx` and contain ERP data, where columns represent channels and rows represent time points.

2. Output:
    - A Word document summarizing the minimum and maximum values of each channel across all conditions and events.
    - A Word document containing detailed descriptive statistics (mean, median, standard deviation, quantiles, etc.) for each condition-event pair.
    - Histogram plots of the ERP data distribution for each condition-event pair, which are embedded in the Word document with statistics.

3. Statistical Methods:
   - Mean: The average value of the ERP data.
   - Median: The middle value of the sorted ERP data.
   - Standard Deviation: A measure of the amount of variation or dispersion of the ERP data.
   - Quantiles: The 25th, 50th, and 75th percentiles of the ERP data.
   - Interquartile Range (IQR): The range between the 25th and 75th percentiles, used to identify outliers.
   - Outliers: Data points that fall below Q1 - 1.5 * IQR or above Q3 + 1.5 * IQR.
   - Minimum and Maximum Values: The smallest and largest values within each channel's ERP data.

4. Procedure:
   - Data Loading: The script reads ERP data from the specified folder structure and organizes it into a dictionary for further processing.
   - Statistics Calculation: For each condition-event pair, the script calculates descriptive statistics for the ERP data across all channels.
   - Min-Max Calculation: It computes the minimum and maximum values for each channel, across all condition-event pairs.
   - Histogram Generation: Histograms of the ERP data distributions are generated for each condition-event pair to visually represent the data spread.
   - Saving Results: The calculated statistics and generated histograms are saved into Word documents for easy review and reporting.

5. Functionality:
   - The script automates the entire process of loading data, calculating statistics, generating visualizations, and saving the results in a report format.
   - It provides warnings if any data files are missing or if data frames are empty.
   - The script is designed to be flexible, allowing it to be adapted to different experimental setups by modifying the conditions and events arrays.

Usage:
- To execute the script, ensure that the folder path is correctly set to the directory containing your ERP data. Running the script will produce the summary documents and visualizations as described above.

Author: Soumya Kaim
Date: 14th August 2024
Version: Final

"""
import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import io

def load_grand_average_data(folder_path):
    data_files = {}
    conditions = ['keypress', 'robot_sync', 'robot_async']
    events = ['64', '65']
    
    for condition in conditions:
        condition_folder = os.path.join(folder_path, condition)
        if os.path.isdir(condition_folder):
            data_files[condition] = {}
            for event in events:
                file_name = f'average_erp_data_{condition}_event_{event}.xlsx'
                file_path = os.path.join(condition_folder, file_name)
                if os.path.isfile(file_path):
                    df = pd.read_excel(file_path)
                    print(f"Loaded data for {condition} event {event}, shape: {df.shape}")
                    if not df.empty:
                        data_files[condition][event] = df
                    else:
                        print(f"Warning: Empty DataFrame for {condition} event {event}")
                else:
                    print(f"File not found: {file_path}")
    
    return data_files

def calculate_statistics_for_condition_event(data, condition, event):
    """
    Calculates and returns various statistics for the given ERP data of a specific condition and event.

    Parameters:
    data (pandas DataFrame): The ERP data to calculate statistics from.
    condition (str): The condition of the ERP data.
    event (str): The event of the ERP data.

    Returns:
    dict or None: A dictionary containing the calculated statistics (mean, median, standard deviation, quantiles, outliers, min value, max value) or None if no ERP data is available.
    """

    erp_data = data.iloc[:, 1:].values.flatten()  # Flattening the data
    print(f"ERP data from {condition} event {event}, size: {erp_data.size}")
    
    if erp_data.size == 0:
        print(f"No ERP data available for {condition} event {event}.")
        return None
    
    mean_value = np.mean(erp_data)
    median_value = np.median(erp_data)
    std_dev = np.std(erp_data)
    quantiles = np.percentile(erp_data, [25, 50, 75])
    
    Q1, Q3 = np.percentile(erp_data, [25, 75])
    IQR = Q3 - Q1
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR
    outliers = erp_data[(erp_data < lower_bound) | (erp_data > upper_bound)]
    
    stats = {
        'mean': mean_value,
        'median': median_value,
        'std_dev': std_dev,
        'quantiles': quantiles,
        'outliers': outliers,
        'min_value': np.min(erp_data),
        'max_value': np.max(erp_data),
    }
    
    return stats

def calculate_channel_min_max(data_files):
    """
    Calculate the minimum and maximum values for each channel in the given data files.

    Parameters:
        data_files (dict): A dictionary containing data files for different conditions and events.

    Returns:
        dict: A dictionary containing the minimum and maximum values for each channel, organized by condition and event.
    """

    channel_min_max = {}
    for condition, events in data_files.items():
        channel_min_max[condition] = {}
        for event, data in events.items():
            channel_min_max[condition][event] = {}
            for channel in data.columns[1:]:  # Skip the 'time' column
                channel_data = data[channel].values
                channel_min_max[condition][event][channel] = (np.min(channel_data), np.max(channel_data))
    return channel_min_max

def save_channel_min_max_to_word(channel_min_max, output_file):
    """
    Saves the minimum and maximum values of each channel for a given set of conditions and events to a Word document.

    Parameters:
    - channel_min_max (dict): A nested dictionary containing the minimum and maximum values of each channel for each condition and event.
    - output_file (str): The path to the output Word document.

    Returns:
    - None

    This function takes a nested dictionary `channel_min_max` that contains the minimum and maximum values of each channel for each condition and event. It creates a Word document using the `Document` class from the `python-docx` library and adds a heading for the channel min and max values. It then iterates over each condition and event in the `channel_min_max` dictionary and adds a heading for the condition and event. For each channel in the event, it adds a paragraph with the channel name and its minimum and maximum values. Finally, it saves the Word document to the specified `output_file` and prints a message indicating where the channel min and max values were saved.
    """

    doc = Document()
    doc.add_heading('Channel Min and Max Values', 0)

    for condition, events in channel_min_max.items():
        for event, channels in events.items():
            doc.add_heading(f'Condition: {condition}, Event: {event}', level=1)
            for channel, (min_val, max_val) in channels.items():
                doc.add_paragraph(f"{channel}: Min = {min_val:.2f}, Max = {max_val:.2f}")
            doc.add_paragraph()  # Add a blank line between sections

    doc.save(output_file)
    print(f"Channel min and max values saved to {output_file}")


def plot_histogram(erp_data, condition, event, save_dir):
    """
    Plots a histogram of ERP data for a given condition and event, and saves the plot to a specified directory.

    Parameters:
        erp_data (array-like): The ERP data to be plotted.
        condition (str): The condition under which the ERP data was collected.
        event (str): The event for which the ERP data was collected.
        save_dir (str): The directory where the histogram plot will be saved.

    Returns:
        str: The path to the saved histogram plot.
    """

    plt.figure(figsize=(10, 6))
    plt.hist(erp_data, bins=50, color='blue', alpha=0.7)
    plt.xlabel('Amplitude (µV)')
    plt.ylabel('Frequency')
    plt.title(f'Distribution of ERP Data for {condition}, Event {event}')
    
    # Save the histogram plot to a specified directory
    if not os.path.exists(save_dir):
        os.makedirs(save_dir)
    
    histogram_path = os.path.join(save_dir, f'histogram_{condition}_event_{event}.png')
    plt.savefig(histogram_path)
    plt.close()
    
    return histogram_path

def save_statistics_and_plots_to_word(stats_data, output_file, histogram_paths):
    """
    Saves the ERP data statistical summary and corresponding histogram plots to a Word document.

    Parameters:
    - stats_data (dict): A nested dictionary containing the statistical data for each condition and event.
    - output_file (str): The path to the output Word document.
    - histogram_paths (dict): A nested dictionary containing the file paths to the histogram plots for each condition and event.

    Returns:
    - None
    """

    doc = Document()
    doc.add_heading('ERP Data Statistical Summary', 0)

    for condition, events in stats_data.items():
        for event, stats in events.items():
            doc.add_heading(f'Condition: {condition}, Event: {event}', level=1)
            doc.add_paragraph(f"Mean: {stats['mean']:.2f}")
            doc.add_paragraph(f"Median: {stats['median']:.2f}")
            doc.add_paragraph(f"Standard Deviation: {stats['std_dev']:.2f}")
            doc.add_paragraph(f"25th percentile: {stats['quantiles'][0]:.2f}")
            doc.add_paragraph(f"50th percentile: {stats['quantiles'][1]:.2f}")
            doc.add_paragraph(f"75th percentile: {stats['quantiles'][2]:.2f}")
            doc.add_paragraph(f"Min Value: {stats['min_value']:.2f}")
            doc.add_paragraph(f"Max Value: {stats['max_value']:.2f}")
            if len(stats['outliers']) > 0:
                doc.add_paragraph(f"Outliers: {stats['outliers']}")
            else:
                doc.add_paragraph("No outliers detected.")
            doc.add_paragraph()  # Add a blank line between sections
            
            # Add the histogram plot to the document
            histogram_path = histogram_paths[condition][event]
            doc.add_picture(histogram_path, width=Inches(6))
    
    doc.save(output_file)
    print(f"Statistics and plots saved to {output_file}")

def main():
    """
    The main function that performs the following steps:
    1. Checks if the specified folder exists, and exits if it does not.
    2. Loads grand average data from the specified folder.
    3. Calculates the minimum and maximum values for each channel in the data.
    4. Saves the channel min and max values to a Word document.
    5. Calculates statistics for each condition and event in the data.
    6. Plots and saves histograms for each condition and event.
    7. Saves the statistics and plots to a Word document.

    Parameters:
    None

    Returns:
    None
    """

    folder_path = r"E:\Academics\MajorThesis\DATA\Robotics\EEG_mainexp+pilot"
    
    if not os.path.exists(folder_path):
        print(f"Directory '{folder_path}' does not exist. Exiting.")
        return
    
    data_files = load_grand_average_data(folder_path)
    
    if not data_files:
        print("No grand average data found. Exiting.")
        return
    
    # Calculate min and max values for each channel
    channel_min_max = calculate_channel_min_max(data_files)
    
    # Save channel min and max values to a Word document
    save_channel_min_max_to_word(channel_min_max, os.path.join(folder_path, 'Channel_Min_Max_Values.docx'))
    
    stats_data = {}
    histogram_paths = {}
    save_dir = os.path.join(folder_path, 'descriptives_histograms')
    
    for condition, events in data_files.items():
        stats_data[condition] = {}
        histogram_paths[condition] = {}
        for event, data in events.items():
            stats = calculate_statistics_for_condition_event(data, condition, event)
            if stats:
                stats_data[condition][event] = stats
                # Plot and save the histogram
                histogram_paths[condition][event] = plot_histogram(data.iloc[:, 1:].values.flatten(), condition, event, save_dir)
    
    # Save statistics and plots to a Word document
    output_file = os.path.join(folder_path, 'ERP_descriptiveStatistics_Histograms_perchannel.docx')
    save_statistics_and_plots_to_word(stats_data, output_file, histogram_paths)

if __name__ == "__main__":
    main()
